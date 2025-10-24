# main.py

import os
import shutil
from pathlib import Path
# from api_key import key
from io import BytesIO
import re
import unicodedata
import time 

# --- Ambiente .env pedir las API
from dotenv import load_dotenv
load_dotenv()

# --- Importaciones de FastAPI (se añaden Response y StreamingResponse) ---
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

# --- Importaciones de LlamaIndex y Agentes ---
from llama_index.core import Settings
from llama_index.llms.openai import OpenAI
from llama_index.llms.gemini import Gemini 
from llama_index.embeddings.openai import OpenAIEmbedding
from agents.agent1 import run_agent_1
from agents.agent2 import run_agent_2
from agents.agent3 import run_agent_3

# --- Importación para crear el Word ---
from docx import Document
from docx.shared import Pt

# --- 1. CONFIGURACIÓN INICIAL ---  
# os.environ["OPENAI_API_KEY"] = key
# Settings.llm = OpenAI(model="gpt-5-mini", temperature=0)
Settings.embed_model = OpenAIEmbedding(model="text-embedding-3-small")

app = FastAPI(title="Agente de Planificación Curricular API", version="1.0.0")

# --- CONFIGURACIÓN DE CORS ---
origins = ["https://v0-educational-chatbot-design-seven.vercel.app", "https://heirless-tabularly-kairi.ngrok-free.dev", "http://localhost:3000", "http://localhost:5173"]
app.add_middleware(CORSMiddleware, allow_origins=origins, allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

DATA_PATH = Path("./data/")
RESULTADOS_PATH = Path("./resultados/")

# --- 2. FUNCIONES AUXILIARES ---

def sanitize_filename(filename: str) -> str:
    """
    Limpia un nombre de archivo para que sea seguro para las cabeceras HTTP,
    reemplazando caracteres especiales con sus equivalentes ASCII.
    """
    # Normalizar a NFD para separar caracteres base de sus acentos
    nfkd_form = unicodedata.normalize('NFD', filename)
    # Codificar a ASCII, ignorando los caracteres que no se pueden convertir (los acentos)
    only_ascii = nfkd_form.encode('ASCII', 'ignore')
    # Decodificar de nuevo a un string
    return only_ascii.decode('utf-8')

# def create_word_document_from_text(text_content: str) -> BytesIO:
#     """
#     Toma el texto de la planificación y lo convierte en un documento Word (.docx)
#     en memoria, aplicando formato avanzado para títulos, listas, negritas y tablas.
#     """
#     document = Document()
    
#     # Limpiar separadores y texto conversacional del LLM
#     clean_content = re.sub(r'----%%%%----|^\*\*\*$\n|¡Saludos, futuros ingenieros!.*?\n|Estimados estudiantes,\n', '', text_content, flags=re.MULTILINE)

#     is_in_table = False
#     table_lines = []

#     def add_paragraph_with_bolding(doc, text):
#         """Añade un párrafo manejando texto en negrita con **."""
#         p = doc.add_paragraph()
#         parts = re.split(r'(\*\*.*?\*\*)', text)
#         for part in parts:
#             if part.startswith('**') and part.endswith('**'):
#                 p.add_run(part.strip('*')).bold = True
#             else:
#                 p.add_run(part)
    
#     for line in clean_content.split('\n'):
#         stripped_line = line.strip()

#         # Detección de tablas Markdown
#         if stripped_line.startswith('|') and stripped_line.endswith('|'):
#             is_in_table = True
#             table_lines.append(stripped_line)
#             continue
#         elif is_in_table:
#             # Fin de la tabla, añadirla al documento
#             is_in_table = False
#             p = document.add_paragraph()
#             table_text = "\n".join(table_lines)
#             run = p.add_run(table_text)
#             font = run.font
#             font.name = 'Courier New'
#             font.size = Pt(10)
#             table_lines = []
        
#         if not stripped_line:
#             continue
        
#         # Lógica de formato
#         if stripped_line.startswith('# MATRIZ'):
#             document.add_heading(stripped_line.replace('# ', ''), level=0)
#         elif stripped_line.startswith('##'):
#             document.add_heading(stripped_line.replace('## ', ''), level=2)
#         elif stripped_line.startswith('###'):
#             document.add_heading(stripped_line.replace('### ', ''), level=3)
#         elif re.match(r'^\d\.\s', stripped_line):
#             p = document.add_paragraph()
#             p.add_run(stripped_line).bold = True
#             p.paragraph_format.space_before = Pt(12)
#         elif stripped_line.startswith('-'):
#             add_paragraph_with_bolding(document.add_paragraph(style='List Bullet'), stripped_line.lstrip(' -'))
#         else:
#             add_paragraph_with_bolding(document, stripped_line)

#     file_stream = BytesIO()
#     document.save(file_stream)
#     file_stream.seek(0)
#     return file_stream

def create_word_document_from_text(text_content: str) -> BytesIO:
    """
    Toma el texto de la planificación y lo convierte en un documento Word (.docx)
    en memoria, aplicando formato avanzado.
    """
    document = Document()
    document.add_heading('Matriz de Planificación Curricular Unificada', level=0)
    
    clean_content = re.sub(r'----%%%%----', '', text_content)

    def add_paragraph_with_bolding(p, text):
        """Añade texto a un párrafo existente, manejando negritas."""
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part.strip('*')).bold = True
            else:
                p.add_run(part)

    for line in clean_content.split('\n'):
        stripped_line = line.strip()
        if not stripped_line:
            continue
        
        # Lógica de formato
        if stripped_line.startswith('###'):
            # Para los subtítulos de las unidades (ej. ### **UNIDAD I...**)
            document.add_heading(stripped_line.strip('#* '), level=3)
        elif stripped_line.startswith('##'):
            # Para los subtítulos principales (ej. ## Resultado de aprendizaje)
            document.add_heading(stripped_line.strip('# '), level=2)
        elif re.match(r'^\d\.\s', stripped_line):
            p = document.add_paragraph()
            run = p.add_run(stripped_line)
            run.bold = True
            p.paragraph_format.space_before = Pt(12)
        elif stripped_line.startswith('-'):
            text = stripped_line.lstrip(' -')
            p = document.add_paragraph(style='List Bullet')
            add_paragraph_with_bolding(p, text)
        else:
            p = document.add_paragraph()
            add_paragraph_with_bolding(p, stripped_line)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream


# --- 3. LÓGICA DE LA API (ENDPOINTS) ---

@app.on_event("startup")
def on_startup():
    DATA_PATH.mkdir(exist_ok=True)
    RESULTADOS_PATH.mkdir(exist_ok=True)

@app.get("/", response_model=dict)
async def root():
    return {"status": "API de Planificación Curricular en línea"}

@app.post("/procesar-pdf/")
async def procesar_pdf_endpoint(
    pdf: UploadFile = File(..., description="El archivo PDF a procesar."), 
    consulta: str = Form(None, description="Texto opcional del usuario.")
):
    """
    Recibe un PDF, ejecuta la cadena de agentes y devuelve un documento
    Word (.docx) con la planificación final para su descarga.
    """
    # Inicialización de los modelos de IA
    openai_llm = OpenAI(model="gpt-4o", temperature=0)
    gemini_llm = Gemini(model="models/gemini-2.5-flash-preview-09-2025", api_key=os.getenv("GOOGLE_API_KEY"))

    temp_pdf_path = DATA_PATH / pdf.filename
    try:
        with open(temp_pdf_path, "wb") as buffer:
            shutil.copyfileobj(pdf.file, buffer)
        
        print(f"Archivo '{pdf.filename}' recibido. Iniciando pipeline...")
        
        # Ejecutar la cadena de agentes pasando el modelo correspondiente
        # structured_path, _ = run_agent_1(temp_pdf_path)
        # planificacion_parcial_path = run_agent_2(structured_path)
        # run_agent_3(planificacion_parcial_path)
        structured_path, _ = run_agent_1(temp_pdf_path, llm=gemini_llm)
        planificacion_parcial_path = run_agent_2(structured_path, llm=gemini_llm)
        run_agent_3(planificacion_parcial_path, llm=gemini_llm)
        
        print("Pipeline completada. Generando documento Word...")

        planificacion_final_path = RESULTADOS_PATH / "planificacion_final.txt"
        with open(planificacion_final_path, "r", encoding="utf-8") as f:
            planificacion_content = f.read()

        word_file_stream = create_word_document_from_text(planificacion_content)
        
        sanitized_stem = sanitize_filename(Path(pdf.filename).stem)
        output_filename = f"planificacion_{sanitized_stem}.docx"
        
        headers = {
            'Content-Disposition': f'attachment; filename="{output_filename}"'
        }
        
        return StreamingResponse(
            word_file_stream, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )

    except Exception as e:
        print(f"ERROR: Ocurrió un error al procesar el archivo: {e}")
        raise HTTPException(status_code=500, detail=f"Ocurrió un error en el servidor: {e}")

    finally:
        # --- BLOQUE DE LIMPIEZA SEGURO Y ROBUSTO ---
        if 'temp_pdf_path' in locals() and temp_pdf_path.exists():
            # Forzar el cierre del manejador de archivos si es posible
            if pdf and not pdf.file.closed:
                pdf.file.close()
            
            # Intentar borrar el archivo, pero no fallar si está bloqueado
            try:
                # Una pequeña pausa puede ayudar a que Windows libere el archivo
                time.sleep(1) 
                temp_pdf_path.unlink()
                print(f"Archivo temporal '{temp_pdf_path.name}' eliminado.")
            except PermissionError:
                print(f"ADVERTENCIA: No se pudo eliminar el archivo temporal '{temp_pdf_path.name}' porque está bloqueado por otro proceso. Se eliminará en la siguiente ejecución.")
            except Exception as e:
                print(f"ADVERTENCIA: Ocurrió un error inesperado al eliminar el archivo temporal: {e}")
    

